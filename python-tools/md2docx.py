import markdown  # 导入 markdown 库，用于将 Markdown 文本转换为 HTML
from bs4 import BeautifulSoup  # 导入 BeautifulSoup 库，用于解析 HTML
from docx import Document  # 导入 python-docx 库，用于生成 Word 文档
from docx.shared import Pt  # 导入字体设置类，用于设置字体大小
from docx.oxml.ns import qn  # 导入 qn 函数，用于设置字体
import argparse  # 导入 argparse 库，用于解析命令行参数
import os  # 导入 os 库，用于处理文件路径


def set_chinese_font(run, size=12):
    """
    设置 run 对象的字体为中文字体，以确保正确显示
    """
    print(f"设置中文字体: {run.text}, 大小: {size}")
    run.font.name = 'SimSun'  # 将字体设置为宋体（SimSun），可以根据需要更改为其他已安装的中文字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(size)


def process_inline_elements(element, paragraph):
    """
    处理段落中的内联元素，例如加粗、斜体和代码
    """
    print(f"处理段落中的内联元素: {paragraph.text}")
    for child in element.contents:
        print(f"处理子元素: {child}")
        if isinstance(child, str):
            run = paragraph.add_run(child)
            set_chinese_font(run)
        elif child.name in ['strong', 'b']:
            run = paragraph.add_run(child.get_text())
            run.bold = True
            set_chinese_font(run)
        elif child.name in ['em', 'i']:
            run = paragraph.add_run(child.get_text())
            run.italic = True
            set_chinese_font(run)
        elif child.name == 'code':
            run = paragraph.add_run(child.get_text())
            set_code_font(run)
        else:
            # 递归处理其他元素
            process_inline_elements(child, paragraph)


def process_element(element, document):
    """
    处理每个 HTML 元素并将其添加到 Word 文档中
    """
    print(f"处理 HTML 元素: {element.name}")
    if element.name == 'h1':
        paragraph = document.add_heading(level=1)
        run = paragraph.add_run(element.get_text())
        set_chinese_font(run, size=16)
    elif element.name == 'h2':
        paragraph = document.add_heading(level=2)
        run = paragraph.add_run(element.get_text())
        set_chinese_font(run, size=14)
    elif element.name == 'h3':
        paragraph = document.add_heading(level=3)
        run = paragraph.add_run(element.get_text())
        set_chinese_font(run, size=13)
    elif element.name == 'p':
        paragraph = document.add_paragraph()
        process_inline_elements(element, paragraph)
    elif element.name == 'ul':
        for li in element.find_all('li', recursive=False):
            print(f"处理列表项: {li.get_text()}")
            paragraph = document.add_paragraph(style='List Bullet')
            process_inline_elements(li, paragraph)
    elif element.name == 'ol':
        for li in element.find_all('li', recursive=False):
            print(f"处理有序列表项: {li.get_text()}")
            paragraph = document.add_paragraph(style='List Number')
            process_inline_elements(li, paragraph)
    elif element.name == 'blockquote':
        print(f"处理引用块: {element.get_text()}")
        paragraph = document.add_paragraph(style='Quote')
        process_inline_elements(element, paragraph)
    elif element.name == 'pre':
        # 处理代码块
        code_block = element.find('code')
        if code_block:
            print(f"处理代码块: {code_block.get_text()}")
            paragraph = document.add_paragraph()
            run = paragraph.add_run(code_block.get_text())
            set_code_font(run)


def set_code_font(run):
    print(f"设置代码字体: {run.text}")
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    run.font.size = Pt(11)


def main():
    # 解析命令行参数
    parser = argparse.ArgumentParser(description='将 Markdown 文件或文件夹转换为 Word 文档')
    parser.add_argument('input_path', help='输入的 Markdown 文件路径或文件夹路径')
    args = parser.parse_args()

    input_path = args.input_path

    if os.path.isfile(input_path) and input_path.endswith('.md'):
        convert_markdown_to_docx(input_path)
    elif os.path.isdir(input_path):
        for root, _, files in os.walk(input_path):
            for file in files:
                if file.endswith('.md'):
                    input_file = os.path.join(root, file)
                    convert_markdown_to_docx(input_file)
    else:
        print(f"错误: 输入路径 '{input_path}' 不是有效的 Markdown 文件或文件夹。")


def convert_markdown_to_docx(input_file):
    output_file = os.path.splitext(input_file)[0] + '.docx'

    try:
        print(f"读取 Markdown 文件: {input_file}")
        with open(input_file, 'r', encoding='utf-8') as file:
            markdown_text = file.read()
    except FileNotFoundError:
        print(f"错误: 文件 '{input_file}' 未找到。")
        return
    except Exception as error:
        print(f"读取文件 '{input_file}' 时出错: {error}")
        return

    # 将 Markdown 转换为 HTML
    extensions = ['fenced_code', 'tables']
    print("将 Markdown 转换为 HTML")
    html_text = markdown.markdown(markdown_text, extensions=extensions)

    # 解析 HTML 内容
    print("解析 HTML 内容")
    soup = BeautifulSoup(html_text, 'html.parser')

    # 创建一个新的 Word 文档
    print("创建新的 Word 文档")
    document = Document()

    # 处理 HTML 中的每个元素
    for element in soup.contents:
        if isinstance(element, str) and not element.strip():
            continue  # 跳过根目录下的空字符串
        else:
            process_element(element, document)

    # 保存 Word 文档
    save_document(document, output_file)


def save_document(document, output_file):
    try:
        print(f"保存 Word 文档到: {output_file}")
        document.save(output_file)
        print(f"成功保存 Word 文档为 '{output_file}'")
    except Exception as error:
        print(f"保存文件 '{output_file}' 时出错: {error}")


if __name__ == '__main__':
    main()