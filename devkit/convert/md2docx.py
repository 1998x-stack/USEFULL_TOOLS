"""Markdown to DOCX converter."""

import os
from typing import Optional


def markdown_to_docx(input_path: str, output_path: Optional[str] = None) -> str:
    """Convert a Markdown file to DOCX.

    Requires: markdown, python-docx, beautifulsoup4 (pip install devkit-tools[convert])
    """
    try:
        import markdown
        from bs4 import BeautifulSoup
        from docx import Document
    except ImportError:
        raise ImportError("markdown, python-docx, beautifulsoup4 required: pip install devkit-tools[convert]")

    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + ".docx"

    with open(input_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    html = markdown.markdown(md_text, extensions=["fenced_code", "tables"])
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()

    for elem in soup.children:
        if not hasattr(elem, "name") or elem.name is None:
            continue
        text = elem.get_text()
        if elem.name == "h1":
            doc.add_heading(text, level=1)
        elif elem.name == "h2":
            doc.add_heading(text, level=2)
        elif elem.name == "h3":
            doc.add_heading(text, level=3)
        elif elem.name in ("p", "div"):
            doc.add_paragraph(text)
        elif elem.name in ("ul", "ol"):
            for li in elem.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(), style="List Bullet")
        elif elem.name == "pre":
            doc.add_paragraph(text, style="Normal")

    doc.save(output_path)
    return output_path
